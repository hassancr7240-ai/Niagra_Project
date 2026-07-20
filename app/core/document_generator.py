from __future__ import annotations

import io
import os
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Optional

# ── ReportLab PDF ─────────────────────────────────────────────────────────────
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    HRFlowable,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

# ── python-docx ───────────────────────────────────────────────────────────────
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches, Cm

# ── openpyxl ──────────────────────────────────────────────────────────────────
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from app.db.models import Task

# ─── Colour palette ──────────────────────────────────────────────────────────

C_HEADER_BG = colors.HexColor("#1B3A6B")   # dark navy header
C_HEADER_FG = colors.white
C_SAFETY = colors.HexColor("#C0392B")       # red  — safety tasks
C_RUNNING = colors.HexColor("#27AE60")      # green — RUNNING state
C_STOPPED = colors.HexColor("#F39C12")      # amber — STOPPED state
C_POWERED = colors.HexColor("#8E44AD")      # purple — POWERED_OFF state
C_ROW_ALT = colors.HexColor("#F8F9FA")      # light grey alternating row
C_WARN_BG = colors.HexColor("#FDEBD0")      # light orange GMP warning bg
C_TABLE_HEADER = colors.HexColor("#2C3E50") # table column headers
C_BORDER = colors.HexColor("#BDC3C7")


@dataclass
class PMDocument:
    machine_name: str
    machine_id: str
    interval_hours: int
    interval_label: str
    work_order: str
    technician_name: str
    tasks: list[Task]
    parts: list[dict]          # [{"part_number": ..., "description": ...}]
    generated_at: datetime
    watermark_text: str = ""   # e.g. "DRAFT" or timestamp
    notes: str = ""


# ─── Shared styles ─────────────────────────────────────────────────────────────

def _build_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        "PMTitle", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=14,
        textColor=C_HEADER_FG, alignment=TA_CENTER, spaceAfter=2,
    ))
    styles.add(ParagraphStyle(
        "PMSubTitle", parent=styles["Normal"],
        fontName="Helvetica", fontSize=9,
        textColor=C_HEADER_FG, alignment=TA_CENTER, spaceAfter=0,
    ))
    styles.add(ParagraphStyle(
        "FieldLabel", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=8,
        textColor=colors.HexColor("#555555"),
    ))
    styles.add(ParagraphStyle(
        "FieldValue", parent=styles["Normal"],
        fontName="Helvetica", fontSize=8,
    ))
    styles.add(ParagraphStyle(
        "SectionHeader", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=9,
        textColor=C_HEADER_FG, alignment=TA_LEFT, spaceAfter=0,
    ))
    styles.add(ParagraphStyle(
        "TaskDesc", parent=styles["Normal"],
        fontName="Helvetica", fontSize=7.5, leading=10,
    ))
    styles.add(ParagraphStyle(
        "TaskDescSafety", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=7.5, leading=10,
        textColor=C_SAFETY,
    ))
    styles.add(ParagraphStyle(
        "Footer", parent=styles["Normal"],
        fontName="Helvetica-Bold", fontSize=8,
        textColor=colors.white, alignment=TA_CENTER,
    ))
    styles.add(ParagraphStyle(
        "FooterWarn", parent=styles["Normal"],
        fontName="Helvetica", fontSize=7,
        textColor=colors.HexColor("#333333"), alignment=TA_LEFT,
    ))
    return styles


_STATE_LABELS = {
    "RUNNING": "RUNNING",
    "STOPPED": "STOPPED",
    "POWERED_OFF": "POWERED OFF",
}
_STATE_COLORS = {
    "RUNNING": C_RUNNING,
    "STOPPED": C_STOPPED,
    "POWERED_OFF": C_POWERED,
}


def _state_cell(state: str, styles) -> Paragraph:
    colour = _STATE_COLORS.get(state, colors.grey)
    label = _STATE_LABELS.get(state, state)
    style = ParagraphStyle(
        f"state_{state}",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=6.5,
        textColor=colour,
        alignment=TA_CENTER,
    )
    return Paragraph(label, style)


# ─── PDF generator ─────────────────────────────────────────────────────────────

def generate_pdf(doc: PMDocument, output_path: Path) -> Path:
    """Generate a GMP-compliant PM checklist PDF matching the CON L3 template format."""
    styles = _build_styles()
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    story = []
    page_w, page_h = A4
    usable_w = page_w - 2 * cm

    # ── Header block ──────────────────────────────────────────────────────────
    header_data = [
        [
            Paragraph(f"{doc.machine_name.upper()}", styles["PMTitle"]),
            Paragraph(f"{doc.interval_label.upper()} PREVENTIVE MAINTENANCE", styles["PMTitle"]),
        ]
    ]
    header_table = Table(header_data, colWidths=[usable_w * 0.5, usable_w * 0.5])
    header_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), C_HEADER_BG),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
    ]))
    story.append(header_table)

    # ── Field grid ────────────────────────────────────────────────────────────
    field_style = ParagraphStyle(
        "FS", parent=styles["Normal"], fontSize=8, fontName="Helvetica"
    )
    label_style = ParagraphStyle(
        "LS", parent=styles["Normal"], fontSize=8, fontName="Helvetica-Bold",
        textColor=colors.HexColor("#2C3E50"),
    )
    blank = "  " + "_" * 28

    fields = [
        [
            Paragraph("Asset Activity:", label_style),
            Paragraph(f"{doc.machine_name} {doc.interval_label} PM", field_style),
            Paragraph("Work Order #:", label_style),
            Paragraph(doc.work_order, field_style),
        ],
        [
            Paragraph("Machine ID:", label_style),
            Paragraph(doc.machine_id, field_style),
            Paragraph("Actual Start Date:", label_style),
            Paragraph(blank, field_style),
        ],
        [
            Paragraph("Interval:", label_style),
            Paragraph(f"{doc.interval_hours} Operating Hours", field_style),
            Paragraph("Actual End Date:", label_style),
            Paragraph(blank, field_style),
        ],
        [
            Paragraph("Team Member(s):", label_style),
            Paragraph(doc.technician_name, field_style),
            Paragraph("Machine Hours:", label_style),
            Paragraph(blank, field_style),
        ],
        [
            Paragraph("Generated:", label_style),
            Paragraph(doc.generated_at.strftime("%d %b %Y %H:%M UTC"), field_style),
            Paragraph("Comments:", label_style),
            Paragraph(blank, field_style),
        ],
    ]
    col_w = [usable_w * 0.18, usable_w * 0.32, usable_w * 0.18, usable_w * 0.32]
    field_table = Table(fields, colWidths=col_w)
    field_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#ECF0F1")),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("GRID", (0, 0), (-1, -1), 0.25, C_BORDER),
    ]))
    story.append(field_table)
    story.append(Spacer(1, 4 * mm))

    # ── Group tasks by machine state ──────────────────────────────────────────
    state_order = ["RUNNING", "STOPPED", "POWERED_OFF"]
    state_groups: dict[str, list[Task]] = {s: [] for s in state_order}
    for t in sorted(doc.tasks, key=lambda x: x.task_no):
        state_groups[t.machine_state].append(t)

    for state in state_order:
        group = state_groups[state]
        if not group:
            continue

        # State section header
        state_label = _STATE_LABELS[state]
        state_colour = _STATE_COLORS[state]
        sec_data = [[Paragraph(f"  {state_label} TASKS", styles["SectionHeader"])]]
        sec_t = Table(sec_data, colWidths=[usable_w])
        sec_t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), state_colour),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(sec_t)

        # Task table header
        th_style = ParagraphStyle(
            "TH", parent=styles["Normal"],
            fontName="Helvetica-Bold", fontSize=7,
            textColor=colors.white, alignment=TA_CENTER,
        )
        header_row = [
            Paragraph("TASK #", th_style),
            Paragraph("AREA", th_style),
            Paragraph("ACTION", th_style),
            Paragraph("DESCRIPTION", th_style),
            Paragraph("INITIAL", th_style),
            Paragraph("DONE ☐", th_style),
        ]
        col_widths = [
            usable_w * 0.06,
            usable_w * 0.09,
            usable_w * 0.09,
            usable_w * 0.56,
            usable_w * 0.10,
            usable_w * 0.10,
        ]
        rows = [header_row]

        for i, task in enumerate(group):
            bg = C_ROW_ALT if i % 2 == 0 else colors.white
            if task.safety_flag:
                bg = colors.HexColor("#FADBD8")  # light red for safety

            desc_text = task.description
            if task.part_number:
                desc_text += f"  [PART: {task.part_number}]"

            desc_style = styles["TaskDescSafety"] if task.safety_flag else styles["TaskDesc"]

            task_no_style = ParagraphStyle(
                "TN", parent=styles["Normal"],
                fontName="Helvetica-Bold", fontSize=8, alignment=TA_CENTER,
            )
            action_style = ParagraphStyle(
                "AC", parent=styles["Normal"],
                fontName="Helvetica-Bold", fontSize=7, alignment=TA_CENTER,
                textColor=C_SAFETY if task.safety_flag else C_TABLE_HEADER,
            )
            area_style = ParagraphStyle(
                "AR", parent=styles["Normal"],
                fontName="Helvetica", fontSize=7, alignment=TA_CENTER,
            )

            row = [
                Paragraph(str(task.task_no), task_no_style),
                Paragraph(task.area, area_style),
                Paragraph(task.action, action_style),
                Paragraph(desc_text, desc_style),
                Paragraph("", styles["Normal"]),
                Paragraph("☐", ParagraphStyle("CB", parent=styles["Normal"], fontSize=10, alignment=TA_CENTER)),
            ]
            rows.append(row)

        task_table = Table(rows, colWidths=col_widths, repeatRows=1)
        ts = TableStyle([
            # Header row
            ("BACKGROUND", (0, 0), (-1, 0), C_TABLE_HEADER),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, 0), 7),
            ("ALIGN", (0, 0), (-1, 0), "CENTER"),
            # All cells
            ("FONTSIZE", (0, 1), (-1, -1), 7.5),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("GRID", (0, 0), (-1, -1), 0.25, C_BORDER),
        ])
        # Row backgrounds
        for i in range(1, len(rows)):
            task = group[i - 1]
            if task.safety_flag:
                ts.add("BACKGROUND", (0, i), (-1, i), colors.HexColor("#FADBD8"))
            elif i % 2 == 0:
                ts.add("BACKGROUND", (0, i), (-1, i), C_ROW_ALT)
        task_table.setStyle(ts)
        story.append(task_table)
        story.append(Spacer(1, 3 * mm))

    # ── Sign-off section ──────────────────────────────────────────────────────
    story.append(Spacer(1, 3 * mm))
    signoff_data = [
        [
            Paragraph("Technician Name:", label_style),
            Paragraph(blank, field_style),
            Paragraph("Supervisor Approval:", label_style),
            Paragraph(blank, field_style),
            Paragraph("Date:", label_style),
            Paragraph(blank, field_style),
        ]
    ]
    so_col_w = [usable_w * 0.15, usable_w * 0.20, usable_w * 0.18, usable_w * 0.20, usable_w * 0.08, usable_w * 0.19]
    signoff_table = Table(signoff_data, colWidths=so_col_w)
    signoff_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#ECF0F1")),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("GRID", (0, 0), (-1, -1), 0.25, C_BORDER),
    ]))
    story.append(signoff_table)

    # ── Parts table ───────────────────────────────────────────────────────────
    if doc.parts:
        story.append(Spacer(1, 4 * mm))
        parts_header_data = [[
            Paragraph("  PARTS / CONSUMABLES REQUIRED", styles["SectionHeader"])
        ]]
        parts_h = Table(parts_header_data, colWidths=[usable_w])
        parts_h.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), C_HEADER_BG),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(parts_h)

        parts_col_head = ParagraphStyle(
            "PCH", parent=styles["Normal"],
            fontName="Helvetica-Bold", fontSize=7,
            textColor=colors.white, alignment=TA_CENTER,
        )
        parts_rows = [[
            Paragraph("PART NUMBER", parts_col_head),
            Paragraph("DESCRIPTION", parts_col_head),
            Paragraph("UOM", parts_col_head),
            Paragraph("REQ QTY", parts_col_head),
            Paragraph("ISSUED QTY", parts_col_head),
            Paragraph("PICK LOCATION", parts_col_head),
            Paragraph("MECHANIC INITIAL", parts_col_head),
        ]]
        parts_cw = [usable_w * 0.15, usable_w * 0.28, usable_w * 0.07, usable_w * 0.10,
                    usable_w * 0.10, usable_w * 0.15, usable_w * 0.15]
        pn_style = ParagraphStyle(
            "PN", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=7
        )
        pd_style = ParagraphStyle(
            "PD", parent=styles["Normal"], fontName="Helvetica", fontSize=7
        )
        for i, p in enumerate(doc.parts):
            bg = C_ROW_ALT if i % 2 == 0 else colors.white
            parts_rows.append([
                Paragraph(p.get("part_number", ""), pn_style),
                Paragraph(p.get("description", ""), pd_style),
                Paragraph("EA", pd_style),
                Paragraph("", pd_style),
                Paragraph("", pd_style),
                Paragraph("", pd_style),
                Paragraph("", pd_style),
            ])
        parts_table = Table(parts_rows, colWidths=parts_cw, repeatRows=1)
        parts_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), C_TABLE_HEADER),
            ("GRID", (0, 0), (-1, -1), 0.25, C_BORDER),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        for i in range(1, len(parts_rows)):
            if i % 2 == 0:
                parts_table._argH[i] = None
                parts_table.setStyle(
                    TableStyle([("BACKGROUND", (0, i), (-1, i), C_ROW_ALT)])
                )
        story.append(parts_table)

    # ── GMP Footer ────────────────────────────────────────────────────────────
    story.append(Spacer(1, 6 * mm))
    end_data = [[Paragraph("**** END OF REPORT ****", styles["Footer"])]]
    end_t = Table(end_data, colWidths=[usable_w])
    end_t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), C_HEADER_BG),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(end_t)

    # Compute content hash for integrity watermark (embedded IN document per architecture spec)
    import hashlib
    content_sig = hashlib.sha256(
        f"{doc.machine_id}|{doc.interval_hours}|{doc.work_order}|{doc.generated_at.isoformat()}".encode()
    ).hexdigest()[:16]

    gmp_lines = [
        "GENERAL — PRIOR TO RETURNING TO SERVICE FOLLOW ALL GMP PROCEDURES",
        "⚠  Remove all LOTO (Lockout/Tagout) devices before returning to service",
        "⚠  Re-check all E-STOP (Emergency Stop) and control switches",
        "⚠  This document must be signed and retained as a maintenance record",
        f"Doc-ID: {content_sig} | WO: {doc.work_order} | Generated: {doc.generated_at.strftime('%Y-%m-%d %H:%M')} UTC | PM Automation System v1.0",
    ]
    gmp_data = [[Paragraph(line, styles["FooterWarn"])] for line in gmp_lines]
    gmp_t = Table(gmp_data, colWidths=[usable_w])
    gmp_t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), C_WARN_BG),
        ("TOPPADDING", (0, 0), (-1, -1), 2),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#E59866")),
    ]))
    story.append(gmp_t)

    # ── Build PDF ─────────────────────────────────────────────────────────────
    def _page_header_footer(canvas, pdf_doc):
        canvas.saveState()
        # Top bar: confidential label + page info + timestamp watermark (architecture: "watermark as timestamp embedded")
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(colors.HexColor("#95A5A6"))
        ts = doc.generated_at.strftime("%Y-%m-%d %H:%M UTC")
        canvas.drawRightString(
            page_w - 1.5 * cm, page_h - 0.7 * cm,
            f"Page {pdf_doc.page} | {ts} | Doc-ID: {content_sig}"
        )
        canvas.drawString(
            1.5 * cm, page_h - 0.7 * cm,
            f"CONFIDENTIAL — {doc.machine_id} | {doc.work_order} | {doc.machine_name}"
        )
        # Bottom watermark
        canvas.setFont("Helvetica", 6)
        canvas.setFillColor(colors.HexColor("#CCCCCC"))
        canvas.drawCentredString(
            page_w / 2, 0.5 * cm,
            f"PM Automation System v1.0 | GMP Compliant | Hash: {content_sig} | {ts}"
        )
        canvas.restoreState()

    pdf = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=1 * cm,
        rightMargin=1 * cm,
        topMargin=1.2 * cm,
        bottomMargin=1 * cm,
        title=f"{doc.machine_name} {doc.interval_label} PM",
        author="PM Automation System",
        subject="Preventive Maintenance Checklist",
        creator="PM Automation v1.0",
    )
    pdf.build(story, onFirstPage=_page_header_footer, onLaterPages=_page_header_footer)
    return output_path


# ─── DOCX generator ────────────────────────────────────────────────────────────

def generate_docx(doc: PMDocument, output_path: Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    d = DocxDocument()

    # Title
    title = d.add_heading(f"{doc.machine_name.upper()} — {doc.interval_label} PM", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Header fields
    d.add_paragraph(f"Work Order: {doc.work_order} | Technician: {doc.technician_name}")
    d.add_paragraph(f"Generated: {doc.generated_at.strftime('%d %b %Y %H:%M UTC')}")
    d.add_paragraph(f"Actual Start Date: {'_' * 20}  Actual End Date: {'_' * 20}")
    d.add_paragraph("")

    # Task table
    headers = ["Task #", "Area", "Action", "Description", "Initial", "Done"]
    table = d.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True

    state_order = ["RUNNING", "STOPPED", "POWERED_OFF"]
    for task in sorted(doc.tasks, key=lambda x: (state_order.index(x.machine_state), x.task_no)):
        row_cells = table.add_row().cells
        row_cells[0].text = str(task.task_no)
        row_cells[1].text = task.area
        row_cells[2].text = task.action
        desc = task.description
        if task.part_number:
            desc += f" [PART: {task.part_number}]"
        row_cells[3].text = desc
        row_cells[4].text = ""
        row_cells[5].text = "☐"

    d.add_paragraph("")
    d.add_paragraph("**** END OF REPORT ****")
    d.add_paragraph("PRIOR TO RETURNING TO SERVICE FOLLOW ALL GMP PROCEDURES")
    d.add_paragraph("Remove all LOTO devices before returning to service")
    d.add_paragraph("Re-check all E-STOP and control switches")

    d.save(str(output_path))
    return output_path


# ─── XLSX generator ───────────────────────────────────────────────────────────
# Matches the Niagara/Krones "Asset Activity" Excel template exactly:
#   Separate tab per PM interval (e.g. "500hr (Monthly)", "6000hr (Annual)")
#   Each tab: Row 1 Asset Activity header, Row 3 column headers, tasks, parts,
#   sign-off fields, END OF REPORT footer

_INTERVAL_LABELS = {
    8: "8hr (Daily)", 100: "100hr (2-Week)", 120: "120hr (2-Week)",
    240: "240hr (Monthly)", 500: "500hr (Monthly)", 1000: "1000hr (2-Month)",
    1500: "1500hr (Quarterly)", 3000: "3000hr (6-Month)", 4000: "4000hr (6-Month)",
    6000: "6000hr (Annual)", 12000: "12000hr (2-Year)", 18000: "18000hr (3-Year)",
    30000: "30000hr (5-Year)", 42000: "42000hr (7-Year)", 45000: "45000hr",
}


def _xlsx_interval_label(hours: int) -> str:
    return _INTERVAL_LABELS.get(hours, f"{hours}hr")


# File-safe label for filenames (no spaces/parens)
_INTERVAL_FILE_LABELS = {
    8: "8hr_Daily", 100: "100hr_2-Week", 120: "120hr_2-Week",
    240: "240hr_Monthly", 500: "500hr_Monthly", 1000: "1000hr_2-Month",
    1500: "1500hr_Quarterly", 3000: "3000hr_6-Month", 4000: "4000hr_6-Month",
    6000: "6000hr_Annual", 12000: "12000hr_2-Year", 18000: "18000hr_3-Year",
    30000: "30000hr_5-Year", 42000: "42000hr_7-Year", 45000: "45000hr",
}


def _xlsx_file_label(hours: int) -> str:
    return _INTERVAL_FILE_LABELS.get(hours, f"{hours}hr")


_THIN = Side(style="thin", color="000000")
_BORDER_ALL = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)
_FONT = Font(name="Calibri", size=11, bold=False, color="000000")

# Exact column widths (pt) taken from the approved reference file's <col> tags
_COL_WIDTHS_PT = [67, 214, 58, 232, 85, 99, 122, 85, 85, 48, 48]  # A..K


def _px_width(pt):
    px = pt * 96 / 72
    return round((px - 5) / 7, 1)


def _set_col_widths(ws):
    for i, pt in enumerate(_COL_WIDTHS_PT, start=1):
        ws.column_dimensions[get_column_letter(i)].width = _px_width(pt)


def _fmt_cell(ws, row, col, value=None, align="left", wrap=True):
    c = ws.cell(row=row, column=col, value=value)
    c.font = _FONT
    c.border = _BORDER_ALL
    c.alignment = Alignment(horizontal=align, vertical="top", wrap_text=wrap)
    return c


def _merge_box(ws, row, col_start, col_end, row_end=None):
    r2 = row_end or row
    ws.merge_cells(start_row=row, start_column=col_start, end_row=r2, end_column=col_end)
    for r in range(row, r2 + 1):
        for c in range(col_start, col_end + 1):
            ws.cell(row=r, column=c).border = _BORDER_ALL
            ws.cell(row=r, column=c).font = _FONT


def _write_xlsx_sheet(ws, machine_name, interval_label, tasks,
                      interval_hours=None, work_order=None,
                      technician_name=None, generated_at=None):
    """
    Renders the CON L3 <MACHINE> <HOURS> HOURS PM format exactly:
      Row 1        : Asset Activity: | title | Description: | title
      Row 2        : blank spacer
      Row 3        : Task | Function/Area(B:D) | Action(E:I) | Initial(J:K)
      Row 4..N     : task_no | area(B:D) | description(E:I) | blank(J:K)
      Row N+1..N+2 : Part Number(A:B) | UOM(C) | Required Qty(D:E) |
                     Issued Qty(F) | Pick Location(G) | blank(H) | Mechanic(I:J)
      footer       : Actual start/end date, Comments, Team Member(s)
      last row     : **** END OF REPORT ****  (centered, full width)
    """
    hours = interval_hours or 0
    _set_col_widths(ws)
    ws.sheet_view.showGridLines = False

    title = f"CON L3 {machine_name.upper()} {hours} HOURS PM"

    # Row 1 -- title bar
    _fmt_cell(ws, 1, 1, "Asset Activity:", align="right")
    _fmt_cell(ws, 1, 2, title, align="left")
    _fmt_cell(ws, 1, 3, "Description:", align="left")
    _fmt_cell(ws, 1, 4, title, align="left")
    for c in range(5, 12):
        ws.cell(row=1, column=c).border = _BORDER_ALL
        ws.cell(row=1, column=c).font = _FONT
    ws.row_dimensions[1].height = 13.2

    # Row 2 blank spacer -> header starts row 3
    row = 3
    _fmt_cell(ws, row, 1, "Task", align="left")
    _merge_box(ws, row, 2, 4)
    ws.cell(row=row, column=2, value="Function / Area")
    _merge_box(ws, row, 5, 9)
    ws.cell(row=row, column=5, value="Action")
    _merge_box(ws, row, 10, 11)
    ws.cell(row=row, column=10, value="Initial")
    for c in (2, 5, 10):
        ws.cell(row=row, column=c).alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
    row += 1

    # Task rows
    for t in tasks:
        task_no = getattr(t, "task_no", None) if not isinstance(t, dict) else t.get("task_no")
        area = getattr(t, "area", None) if not isinstance(t, dict) else t.get("area")
        desc = getattr(t, "description", None) if not isinstance(t, dict) else t.get("description")

        _fmt_cell(ws, row, 1, task_no, align="right")
        _merge_box(ws, row, 2, 4)
        ws.cell(row=row, column=2, value=(area or "").upper())
        ws.cell(row=row, column=2).alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        _merge_box(ws, row, 5, 9)
        ws.cell(row=row, column=5, value=(desc or "").upper())
        ws.cell(row=row, column=5).alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)
        _merge_box(ws, row, 10, 11)
        row += 1

    row += 1  # blank spacer before parts table

    # Parts header (2 rows tall)
    pr = row
    _merge_box(ws, pr, 1, 2, pr + 1)
    ws.cell(row=pr, column=1, value="Part Number")
    ws.cell(row=pr, column=1).alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)

    _merge_box(ws, pr, 3, 3, pr + 1)
    ws.cell(row=pr, column=3, value="UOM")
    ws.cell(row=pr, column=3).alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)

    _merge_box(ws, pr, 4, 5, pr + 1)
    ws.cell(row=pr, column=4, value="Required Qty")
    ws.cell(row=pr, column=4).alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)

    _merge_box(ws, pr, 6, 6, pr + 1)
    ws.cell(row=pr, column=6, value="Issued Qty")
    ws.cell(row=pr, column=6).alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)

    _merge_box(ws, pr, 7, 7, pr + 1)
    ws.cell(row=pr, column=7, value="Pick Location")
    ws.cell(row=pr, column=7).alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)

    _merge_box(ws, pr, 8, 8, pr + 1)  # blank column, 2 rows

    _merge_box(ws, pr, 9, 10)  # Mechanic, row 1 only
    ws.cell(row=pr, column=9, value="Mechanic")
    ws.cell(row=pr, column=9).alignment = Alignment(horizontal="right", vertical="top", wrap_text=True)
    _fmt_cell(ws, pr, 11, None)

    _merge_box(ws, pr + 1, 9, 10)  # row 2 continuation of Mechanic box
    _fmt_cell(ws, pr + 1, 11, None)

    row = pr + 2
    row += 1  # blank spacer (no border)

    footer_lines = [
        "Actual start date: ___________   Actual end date: ___________   Hours ___________",
        "Comments & Action: ____________________________________________________",
        "Team Member(s): _______________________________________________________",
    ]
    for line in footer_lines:
        c = ws.cell(row=row, column=1, value=line)
        c.font = _FONT
        c.alignment = Alignment(horizontal="left", vertical="top")
        row += 1

    row += 1  # blank spacer

    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=11)
    c = ws.cell(row=row, column=1,
                value="**************************** END OF REPORT **************************")
    c.font = _FONT
    c.alignment = Alignment(horizontal="center", vertical="center")


def generate_xlsx(doc: PMDocument, output_path: Path) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb = Workbook()
    wb.remove(wb.active)

    # Group tasks by interval_hours (separate tab per PM frequency)
    from collections import defaultdict
    by_interval: dict[int, list] = defaultdict(list)
    for task in sorted(doc.tasks, key=lambda x: x.task_no):
        iv = getattr(task, "interval_hours", 0) or 0
        by_interval[iv].append(task)

    if not by_interval or (len(by_interval) == 1 and 0 in by_interval):
        # No interval data — single sheet with all tasks
        ws = wb.create_sheet(title=f"{doc.interval_label} PM".upper()[:31])
        _write_xlsx_sheet(ws, doc.machine_name, doc.interval_label, doc.tasks,
                          interval_hours=doc.interval_hours, work_order=doc.work_order,
                          technician_name=doc.technician_name, generated_at=doc.generated_at)
    else:
        for interval in sorted(by_interval.keys()):
            label = _xlsx_interval_label(interval)
            ws = wb.create_sheet(title=label.upper()[:31])
            _write_xlsx_sheet(ws, doc.machine_name, label, by_interval[interval],
                              interval_hours=interval, work_order=doc.work_order,
                              technician_name=doc.technician_name, generated_at=doc.generated_at)

    wb.save(str(output_path))
    return output_path


def _sanitize_con_l3_name(name: str) -> str:
    """Convert a machine name to uppercase CON L3 display format."""
    for src, dst in [('ä', 'a'), ('ö', 'o'), ('ü', 'u'), ('Ä', 'A'), ('Ö', 'O'),
                     ('Ü', 'U'), ('ß', 'ss'), ('.', ''), ('_', ' ')]:
        name = name.replace(src, dst)
    import re
    name = re.sub(r'[^\w\s\-]', '', name)
    return ' '.join(name.upper().split())


# Map machine_id → (con_l3_display_name, zip_stem)
_MACHINE_CON_L3: dict[str, tuple[str, str]] = {
    'DEHUMIDIFIER-L3':   ('EISBAR DAS-E8K2 DEHUMIDIFIER', 'Dehumidifier_EISBAR'),
    'CONTIFORM-C3-L3':   ('KRONES CONTIFORM C3 SAN',      'Krones_Contiform_C3_SAN'),
    'VARIOPAC-PRO-L3':   ('KRONES VARIOPAC PRO FS',        'Variopac_Pro_KRONES'),
    'SHRINK-TUNNEL-L3':  ('KRONES SHRINK TUNNEL',          'Shrink_Tunnel_KRONES'),
    'BOTTLECODER-L3':    ('BOTTLE CODER',                  'CON_L3_Bottle_Coder'),
    'TETRAPAK-ASEPTIC-L3': ('TETRA PAK ASEPTIC TANK',     'Tetra_Pak_Aseptic_Tank'),
}


def generate_con_l3_zip_bytes(machine_id: str, machine_display_name: str, tasks: list) -> tuple[bytes, str]:
    """
    Generate a ZIP in memory with one XLSX per PM interval.
    Returns (zip_bytes, zip_filename).
    """
    import io, zipfile
    from collections import defaultdict

    con_l3_name, zip_stem = _MACHINE_CON_L3.get(machine_id, (None, None))
    if not con_l3_name:
        con_l3_name = _sanitize_con_l3_name(machine_display_name)
    if not zip_stem:
        zip_stem = con_l3_name.title().replace(' ', '_')

    safe_prefix = con_l3_name.replace(' ', '_').replace('-', '_')

    by_interval: dict[int, list] = defaultdict(list)
    for task in sorted(tasks,
                       key=lambda x: getattr(x, 'task_no', 0) if not isinstance(x, dict)
                                     else x.get('task_no', 0)):
        iv = (int(getattr(task, 'interval_hours', 0) or 0)
              if not isinstance(task, dict)
              else int(task.get('interval_hours', 0) or 0))
        by_interval[iv].append(task)

    # If all tasks landed in interval 0 (no interval data) use a single sheet
    if not by_interval or list(by_interval.keys()) == [0]:
        by_interval = {0: tasks}

    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for interval in sorted(by_interval.keys()):
            label = _xlsx_interval_label(interval)
            file_label = _xlsx_file_label(interval)
            xlsx_name = f"PM_{safe_prefix}_{file_label}.xlsx"
            wb = Workbook()
            ws = wb.active
            ws.title = label.upper()[:31]
            _write_xlsx_sheet(ws, con_l3_name, label, by_interval[interval],
                              interval_hours=interval)
            wb_bytes_buf = io.BytesIO()
            wb.save(wb_bytes_buf)
            zf.writestr(xlsx_name, wb_bytes_buf.getvalue())

    return zip_buf.getvalue(), f"{zip_stem}.zip"


def generate_xlsx_all_intervals(doc: PMDocument, output_dir: Path) -> list[tuple[Path, str, int]]:
    """
    Generate one XLSX file per PM interval in output_dir.
    Returns list of (file_path, interval_label, interval_hours) tuples, sorted by interval.
    Each file matches the Niagara Asset Activity template exactly.
    """
    from collections import defaultdict

    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    by_interval: dict[int, list] = defaultdict(list)
    for task in sorted(doc.tasks, key=lambda x: getattr(x, "task_no", 0)):
        iv = getattr(task, "interval_hours", 0) or 0
        by_interval[iv].append(task)

    safe_machine = (
        doc.machine_id.replace(" ", "_").replace("-", "_")
        .replace(".", "").replace("/", "_").upper()
    )

    if not by_interval or (len(by_interval) == 1 and 0 in by_interval):
        label = _xlsx_interval_label(doc.interval_hours or 0)
        file_label = _xlsx_file_label(doc.interval_hours or 0)
        file_name = f"PM_{safe_machine}_{file_label}.xlsx"
        out_path = output_dir / file_name
        wb = Workbook()
        ws = wb.active
        ws.title = label.upper()[:31]
        _write_xlsx_sheet(ws, doc.machine_name, label, doc.tasks,
                          interval_hours=doc.interval_hours or 0, work_order=doc.work_order,
                          technician_name=doc.technician_name, generated_at=doc.generated_at)
        wb.save(str(out_path))
        return [(out_path, label, doc.interval_hours or 0)]

    results = []
    for interval in sorted(by_interval.keys()):
        label = _xlsx_interval_label(interval)
        file_label = _xlsx_file_label(interval)
        file_name = f"PM_{safe_machine}_{file_label}.xlsx"
        out_path = output_dir / file_name
        wb = Workbook()
        ws = wb.active
        ws.title = label.upper()[:31]
        _write_xlsx_sheet(ws, doc.machine_name, label, by_interval[interval],
                          interval_hours=interval, work_order=doc.work_order,
                          technician_name=doc.technician_name, generated_at=doc.generated_at)
        wb.save(str(out_path))
        results.append((out_path, label, interval))

    return results
